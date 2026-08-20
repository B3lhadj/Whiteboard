from flask import Flask, redirect, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash
import base64
import os
import re
import secrets
import smtplib
from email.message import EmailMessage
from io import BytesIO
import shutil
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
import cloudmersive_convert_api_client
from cloudmersive_convert_api_client.rest import ApiException

import fitz

try:
    from pdf2docx import Converter
except Exception:
    Converter = None

PML_NS = 'http://schemas.openxmlformats.org/presentationml/2006/main'
DRAW_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
PKG_REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'


def emu_to_percent(value, total):
    try:
        return (int(value) / int(total)) * 100
    except Exception:
        return 0

app = Flask(__name__)
CORS(app)

ALLOWED_EXTENSIONS = {'pptx'}
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

MONGODB_URI = os.environ.get('MONGODB_URI') or os.environ.get('MONGO_URI') or 'mongodb://localhost:27017'
MONGODB_DB_NAME = os.environ.get('MONGODB_DB') or 'office_editor'
_mongo_db = None
_mongo_error = None


def get_mongo_db():
    """Lazy MongoDB connection so the editor can still boot without MongoDB."""
    global _mongo_db, _mongo_error

    if _mongo_db is not None:
        return _mongo_db

    try:
        from pymongo import MongoClient, DESCENDING

        client = MongoClient(MONGODB_URI, serverSelectionTimeoutMS=1800)
        client.admin.command('ping')
        db = client[MONGODB_DB_NAME]
        db.files.create_index('fileId', unique=True)
        db.files.create_index([('uploadedAt', DESCENDING)])
        db.users.create_index('userId', unique=True)
        db.users.create_index('email', unique=True, sparse=True)
        db.file_shares.create_index([('fileId', 1), ('sharedWith.userId', 1)], unique=True)
        try:
            db.file_shares.create_index('accessToken', unique=True, sparse=True)
        except Exception as index_error:
            if 'IndexKeySpecsConflict' not in str(index_error) and 'same name as the requested index' not in str(index_error):
                raise
            db.file_shares.drop_index('accessToken_1')
            db.file_shares.create_index('accessToken', unique=True, sparse=True)
        db.file_shares.create_index([('sharedWith.userId', 1), ('createdAt', DESCENDING)])
        db.edit_events.create_index([('fileId', 1), ('createdAt', DESCENDING)])
        db.edit_events.create_index([('createdAt', DESCENDING)])
        _mongo_db = db
        _mongo_error = None
        return _mongo_db
    except Exception as error:
        _mongo_error = str(error)
        return None


def serialize_mongo_document(document):
    if isinstance(document, list):
        return [serialize_mongo_document(item) for item in document]
    if isinstance(document, dict):
        serialized = {}
        for key, value in document.items():
            serialized[key] = str(value) if key == '_id' else serialize_mongo_document(value)
        return serialized
    if isinstance(document, datetime):
        return document.isoformat()
    return document


def get_request_user_agent():
    user_agent = request.headers.get('User-Agent', '')
    return user_agent[:240]


def normalize_email(value):
    email = str(value or '').strip().lower()
    if not email:
        return ''
    if not re.match(r'^[^@\s]+@[^@\s]+\.[^@\s]+$', email):
        return ''
    return email[:254]


def get_actor_from_payload(data):
    email = normalize_email(data.get('email') or data.get('userEmail') or data.get('editorEmail'))
    display_name = str(
        data.get('editorName')
        or data.get('displayName')
        or data.get('userName')
        or data.get('userId')
        or email
        or 'Local user'
    ).strip()
    display_name = display_name[:160] or 'Local user'
    user_id = str(data.get('userId') or email or display_name).strip()[:120] or 'local-user'
    actor = {
        'userId': user_id,
        'displayName': display_name,
    }
    if email:
        actor['email'] = email
    return actor


def ensure_user(db, actor):
    now = datetime.now(timezone.utc)
    user_doc = {
        'userId': actor['userId'],
        'displayName': actor['displayName'],
        'updatedAt': now,
        'lastSeenAt': now,
    }
    if actor.get('email'):
        user_doc['email'] = actor['email']

    db.users.update_one(
        {'userId': actor['userId']},
        {
            '$set': user_doc,
            '$setOnInsert': {'createdAt': now},
        },
        upsert=True,
    )
    return user_doc


def get_public_base_url():
    return (os.environ.get('APP_PUBLIC_URL') or os.environ.get('BACKEND_PUBLIC_URL') or request.host_url.rstrip('/')).rstrip('/')


def get_frontend_public_url():
    configured_url = os.environ.get('FRONTEND_PUBLIC_URL') or os.environ.get('APP_FRONTEND_URL')
    if configured_url:
        return configured_url.rstrip('/')

    host = request.host.split(':')[0]
    if host in ['localhost', '127.0.0.1']:
        return f'{request.scheme}://{host}:5173'

    return get_public_base_url()


def save_file_content_to_gridfs(db, file_id, data, existing_file=None):
    content_base64 = data.get('contentBase64')
    if not content_base64:
        return {}

    try:
        file_bytes = base64.b64decode(content_base64)
    except Exception:
        raise ValueError('contentBase64 is not valid base64')

    try:
        import gridfs
        from bson import ObjectId
    except Exception as error:
        raise RuntimeError(f'GridFS is not available: {error}')

    fs = gridfs.GridFS(db)
    old_gridfs_id = (existing_file or {}).get('contentGridFsId')
    if old_gridfs_id:
        try:
            fs.delete(ObjectId(old_gridfs_id))
        except Exception:
            pass

    gridfs_id = fs.put(
        file_bytes,
        filename=str(data.get('fileName') or data.get('name') or file_id),
        content_type=str(data.get('contentType') or 'application/octet-stream'),
        fileId=file_id,
    )
    return {
        'contentStorage': 'gridfs',
        'contentGridFsId': str(gridfs_id),
        'contentType': str(data.get('contentType') or 'application/octet-stream'),
        'contentSize': len(file_bytes),
    }


def send_share_email(to_email, recipient_name, sender_name, file_name, access_url, permission):
    smtp_host = os.environ.get('SMTP_HOST')
    if not smtp_host:
        return {'sent': False, 'configured': False, 'reason': 'SMTP_HOST is not configured'}

    smtp_port = int(os.environ.get('SMTP_PORT') or 587)
    smtp_user = os.environ.get('SMTP_USER')
    smtp_password = os.environ.get('SMTP_PASSWORD')
    smtp_from = os.environ.get('SMTP_FROM') or smtp_user
    use_ssl = os.environ.get('SMTP_SSL', '').lower() == 'true'
    use_tls = os.environ.get('SMTP_TLS', 'true').lower() != 'false'

    if not smtp_from:
        return {'sent': False, 'configured': False, 'reason': 'SMTP_FROM or SMTP_USER is required'}

    message = EmailMessage()
    message['Subject'] = f'{sender_name} shared {file_name} with you'
    message['From'] = smtp_from
    message['To'] = to_email
    message.set_content(
        f'Hello {recipient_name},\n\n'
        f'{sender_name} shared "{file_name}" with you.\n'
        f'Permission: {permission}.\n\n'
        f'Open or download the file here:\n{access_url}\n'
    )

    try:
        smtp_class = smtplib.SMTP_SSL if use_ssl else smtplib.SMTP
        with smtp_class(smtp_host, smtp_port, timeout=10) as server:
            if use_tls and not use_ssl:
                server.starttls()
            if smtp_user and smtp_password:
                server.login(smtp_user, smtp_password)
            server.send_message(message)
        return {'sent': True, 'configured': True}
    except Exception as error:
        return {'sent': False, 'configured': True, 'reason': str(error)}


def build_file_record(data):
    now = datetime.now(timezone.utc)
    file_id = str(data.get('fileId') or '').strip() or str(uuid.uuid4())
    file_name = str(data.get('fileName') or data.get('name') or '').strip()

    if not file_name:
        return None, 'fileName is required'

    actor = get_actor_from_payload(data)
    return {
        'fileId': file_id[:160],
        'fileName': file_name[:260],
        'fileType': str(data.get('fileType') or '')[:40],
        'originalType': str(data.get('originalType') or '')[:40],
        'workflow': str(data.get('workflow') or '')[:80],
        'size': int(data.get('size') or 0),
        'uploadedBy': actor,
        'uploadedAt': now,
        'updatedAt': now,
        'edited': False,
        'editCount': 0,
        'lastEdit': None,
        'ipAddress': request.remote_addr,
        'userAgent': get_request_user_agent(),
    }, None

# Cloudmersive Configuration
CLOUDMERSIVE_API_KEY = 'a0849081-22a5-4eba-aeb1-46c9d394fc47'
configuration = cloudmersive_convert_api_client.Configuration()
configuration.api_key['Apikey'] = CLOUDMERSIVE_API_KEY
api_client = cloudmersive_convert_api_client.ApiClient(configuration)
convert_api = cloudmersive_convert_api_client.ConvertDocumentApi(api_client)

LIBREOFFICE_CANDIDATES = [
    os.environ.get('LIBREOFFICE_PATH', ''),
    os.environ.get('SOFFICE_PATH', ''),
    r'C:\Program Files\LibreOffice\program\soffice.com',
    r'C:\Program Files (x86)\LibreOffice\program\soffice.com',
    shutil.which('soffice') or '',
    shutil.which('libreoffice') or '',
    r'C:\Program Files\LibreOffice\program\soffice.exe',
    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
]

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def find_libreoffice_executable():
    for candidate in LIBREOFFICE_CANDIDATES:
        if candidate and os.path.exists(candidate):
            if candidate.lower().endswith('soffice.exe'):
                com_candidate = candidate[:-4] + '.com'
                if os.path.exists(com_candidate):
                    return com_candidate
            return candidate
    return None


def _color_int_to_rgb(color_value):
    from docx.shared import RGBColor

    try:
        if color_value is None:
            return RGBColor(0, 0, 0)
        if isinstance(color_value, str):
            color_value = color_value.lstrip('#')
            if len(color_value) == 6:
                return RGBColor.from_string(color_value.upper())
        if isinstance(color_value, int):
            hex_value = f'{color_value:06x}'
            return RGBColor.from_string(hex_value.upper())
    except Exception:
        pass
    return RGBColor(0, 0, 0)


def _guess_bold(font_name, flags):
    font_name = (font_name or '').lower()
    return 'bold' in font_name or 'black' in font_name or bool(flags & 16)


def _guess_italic(font_name, flags):
    font_name = (font_name or '').lower()
    return 'italic' in font_name or 'oblique' in font_name or bool(flags & 2)


def _detect_list_bullet(text):
    """Detect if text starts with a bullet or list marker."""
    stripped = text.lstrip()
    bullets = ['•', '○', '◦', '■', '□', '▪', '-', '+', '*']
    for bullet in bullets:
        if stripped.startswith(bullet):
            return True
    import re
    if re.match(r'^[\d]+[\.\)]', stripped):
        return True
    return False


def _is_table_row(lines, page_width):
    """Heuristic to detect if a group of lines forms a table row."""
    if len(lines) < 2:
        return False

    # Check if lines have similar y-coordinates (same row)
    y_values = [line.get('bbox', [0, 0, 0, 0])[1] for line in lines]
    y_range = max(y_values) - min(y_values)
    return y_range < page_width * 0.05


def convert_pdf_to_docx_custom(pdf_path):
    """Enhanced PDF to DOCX conversion with improved formatting, lists, and structure detection."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_LINE_SPACING

    doc = Document()
    pdf = fitz.open(pdf_path)

    last_font_size = 11.0
    consecutive_small_text = 0

    for page_index, page in enumerate(pdf):
        if page_index > 0:
            doc.add_page_break()

        section = doc.sections[-1]
        section.page_width = Pt(page.rect.width)
        section.page_height = Pt(page.rect.height)
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)

        text_dict = page.get_text('dict')
        blocks = text_dict.get('blocks', [])

        # Extract images from the page
        for img_index in range(len(page.get_images())):
            try:
                xref = page.get_images()[img_index]
                pix = fitz.Pixmap(pdf, xref)
                if pix.n - pix.alpha < 4:  # Gray or RGB
                    img_data = pix.tobytes('png')
                else:  # CMYK
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                    img_data = pix.tobytes('png')

                # Add image to document with reasonable sizing
                from io import BytesIO
                img_stream = BytesIO(img_data)
                if len(doc.paragraphs) > 0:
                    last_para = doc.paragraphs[-1]
                    run = last_para.add_run()
                    run.add_picture(img_stream, width=Inches(5.5))
            except Exception as e:
                print(f'Could not extract image from PDF: {e}')

        for block in blocks:
            if block.get('type') != 0:  # Skip non-text blocks
                continue

            lines = block.get('lines', [])
            if not lines:
                continue

            for line_index, line in enumerate(lines):
                spans = line.get('spans', [])
                if not spans:
                    continue

                text_line = ''.join(span.get('text', '') for span in spans).strip()
                if not text_line:
                    continue

                # Create paragraph with enhanced style detection
                paragraph = doc.add_paragraph()

                # Line positioning and alignment
                line_bbox = line.get('bbox', [0, 0, 0, 0])
                page_width = max(page.rect.width, 1)
                left_ratio = line_bbox[0] / page_width
                right_ratio = (page_width - line_bbox[2]) / page_width

                # Alignment detection
                if left_ratio < 0.08 and right_ratio < 0.08:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif right_ratio < 0.12:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Enhanced heading and list detection
                max_size = max(span.get('size', 11) for span in spans)
                last_font_size = max_size

                # Detect headings based on size
                if max_size >= 24:
                    paragraph.style = 'Heading 1'
                    consecutive_small_text = 0
                elif max_size >= 18:
                    paragraph.style = 'Heading 2'
                    consecutive_small_text = 0
                elif max_size >= 14:
                    paragraph.style = 'Heading 3'
                    consecutive_small_text = 0
                elif _detect_list_bullet(text_line):
                    # List detection
                    paragraph.style = 'List Bullet'
                    paragraph.paragraph_format.left_indent = Pt(36)

                # Spacing based on context
                bbox = line.get('bbox', [0, 0, 0, 0])
                y0 = bbox[1] if len(bbox) > 1 else 0
                y1 = bbox[3] if len(bbox) > 3 else 0
                height = max(y1 - y0, 1)

                paragraph.paragraph_format.space_before = Pt(max(min(height * 0.15, 8), 2))
                paragraph.paragraph_format.space_after = Pt(max(min(height * 0.1, 6), 1))
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(height * 1.1)

                # Add text with preserved formatting
                for span in spans:
                    text = span.get('text', '')
                    if not text:
                        continue

                    run = paragraph.add_run(text)
                    font_name = span.get('font', '')
                    flags = int(span.get('flags', 0) or 0)
                    size = float(span.get('size', 11) or 11)

                    # Font styling
                    run.font.size = Pt(size)
                    run.font.bold = _guess_bold(font_name, flags)
                    run.font.italic = _guess_italic(font_name, flags)
                    run.font.color.rgb = _color_int_to_rgb(span.get('color'))

                    # Subscript/superscript detection based on rise
                    rise = span.get('origin', [0, 0])[1] if span.get('origin') else 0
                    if rise > height * 0.3:
                        run.font.superscript = True
                    elif rise < -height * 0.3:
                        run.font.subscript = True

                    if font_name:
                        try:
                            run.font.name = font_name
                        except Exception:
                            pass  # Some font names may not be available

    # Clean up empty paragraphs
    for paragraph in doc.paragraphs[:]:
        if not paragraph.text.strip():
            p = paragraph._element
            p.getparent().remove(p)

    output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    output.close()
    try:
        doc.save(output.name)
        with open(output.name, 'rb') as f:
            docx_bytes = f.read()
        return base64.b64encode(docx_bytes).decode('utf-8')
    finally:
        pdf.close()
        try:
            os.remove(output.name)
        except Exception:
            pass


def convert_pptx_to_png_slides(pptx_path):
    """Convert PPTX to PNG images using LibreOffice -> PDF -> PNG pipeline."""
    libreoffice = find_libreoffice_executable()
    if not libreoffice:
        raise RuntimeError(
            'LibreOffice is not installed or not configured. Set LIBREOFFICE_PATH to soffice.exe.'
        )

    base_name = Path(pptx_path).stem
    work_dir = tempfile.mkdtemp(prefix='pptx_render_')
    pdf_path = os.path.join(work_dir, f'{base_name}.pdf')

    try:
        convert_cmd = [
            libreoffice,
            '--headless',
            '--nologo',
            '--nofirststartwizard',
            '--convert-to', 'pdf',
            '--outdir', work_dir,
            pptx_path,
        ]
        result = subprocess.run(convert_cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(result.stderr.strip() or result.stdout.strip() or 'LibreOffice conversion failed')

        if not os.path.exists(pdf_path):
            # LibreOffice sometimes keeps the source filename; search for any pdf in output dir.
            pdf_candidates = list(Path(work_dir).glob('*.pdf'))
            if not pdf_candidates:
                raise RuntimeError('LibreOffice did not generate a PDF file.')
            pdf_path = str(pdf_candidates[0])

        slides = []
        pdf_doc = fitz.open(pdf_path)
        for page_index, page in enumerate(pdf_doc):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            png_bytes = pixmap.tobytes('png')
            png_base64 = base64.b64encode(png_bytes).decode('utf-8')
            slides.append({
                'id': f'slide-{page_index + 1}',
                'pageNumber': page_index + 1,
                'title': f'Slide {page_index + 1}',
                'imageData': f'data:image/png;base64,{png_base64}',
                'thumbnailData': f'data:image/png;base64,{png_base64}',
                'width': int(page.rect.width),
                'height': int(page.rect.height),
            })
        pdf_doc.close()
        return slides
    finally:
        try:
            shutil.rmtree(work_dir, ignore_errors=True)
        except Exception:
            pass


def convert_pdf_to_docx(pdf_path):
    """Convert PDF to DOCX using Cloudmersive (primary), or custom fitz/pdf2docx logic (fallbacks)."""
    try:
        # Try Cloudmersive first for highest fidelity
        print(f"Attempting Cloudmersive conversion for {pdf_path}")
        api_response = convert_api.convert_document_pdf_to_docx(pdf_path)
        # Cloudmersive returns the file path to the result in some versions,
        # or the bytes directly. The Python client typically returns the bytes.
        if isinstance(api_response, bytes):
            return base64.b64encode(api_response).decode('utf-8')
        elif isinstance(api_response, str) and os.path.exists(api_response):
            with open(api_response, 'rb') as f:
                return base64.b64encode(f.read()).decode('utf-8')
    except Exception as cm_error:
        print(f"Cloudmersive conversion failed: {cm_error}")

    try:
        return convert_pdf_to_docx_custom(pdf_path)
    except Exception as custom_error:
        print(f'Custom PDF->DOCX conversion failed, falling back to pdf2docx/LibreOffice: {custom_error}')

    work_dir = tempfile.mkdtemp(prefix='pdf_docx_')
    docx_path = os.path.join(work_dir, f'{Path(pdf_path).stem}.docx')

    fidelity_settings = {
        'ocr': 0,
        'ignore_page_error': True,
        'multi_processing': False,
        'cpu_count': 1,
        'min_section_height': 12.0,
        'page_margin_factor_top': 0.18,
        'page_margin_factor_bottom': 0.18,
        'shape_min_dimension': 2.0,
        'max_line_spacing_ratio': 1.35,
        'line_overlap_threshold': 0.92,
        'line_break_width_ratio': 0.68,
        'line_break_free_space_ratio': 0.08,
        'line_separate_threshold': 4.0,
        'new_paragraph_free_space_ratio': 0.72,
        'lines_left_aligned_threshold': 1.0,
        'lines_right_aligned_threshold': 1.0,
        'lines_center_aligned_threshold': 2.0,
        'clip_image_res_ratio': 4.0,
        'extract_stream_table': False,
        'parse_lattice_table': True,
        'parse_stream_table': True,
        'delete_end_line_hyphen': True,
        'raw_exceptions': False,
        'list_not_table': True,
    }

    # Higher-fidelity conversion path for selectable-text PDFs.
    if Converter is not None:
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None, **fidelity_settings)
            cv.close()
            with open(docx_path, 'rb') as f:
                docx_bytes = f.read()
            return base64.b64encode(docx_bytes).decode('utf-8')
        except Exception as pdf2docx_error:
            print(f'pdf2docx conversion failed, falling back to LibreOffice: {pdf2docx_error}')

    # Fallback path.
    libreoffice = find_libreoffice_executable()
    if not libreoffice:
        raise RuntimeError('LibreOffice is not installed or not configured. Set LIBREOFFICE_PATH to soffice.exe.')

    try:
        convert_cmd = [
            libreoffice,
            '--headless',
            '--nologo',
            '--nofirststartwizard',
            '--convert-to', 'docx:MS Word 2007 XML',
            '--infilter=writer_pdf_import',
            '--outdir', work_dir,
            pdf_path,
        ]
        result = subprocess.run(convert_cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(result.stderr.strip() or result.stdout.strip() or 'LibreOffice PDF->DOCX conversion failed')

        candidates = list(Path(work_dir).glob('*.docx'))
        if not candidates:
            raise RuntimeError('LibreOffice did not generate a DOCX file from PDF.')

        docx_path = str(candidates[0])
        with open(docx_path, 'rb') as f:
            docx_bytes = f.read()

        return base64.b64encode(docx_bytes).decode('utf-8')
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

def extract_text_from_xml(text_body_elem):
    """Extract formatted text from PowerPoint text body element"""
    runs = []
    for paragraph in text_body_elem.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}p'):
        for text_run in paragraph.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}r'):
            text_elem = text_run.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}t')
            if text_elem is not None:
                text_content = text_elem.text or ''

                # Extract formatting
                run_props = text_run.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}rPr')
                bold = False
                italic = False
                color = None
                font_size = None

                if run_props is not None:
                    bold = run_props.get('b') == '1'
                    italic = run_props.get('i') == '1'

                    # Font size in hundredths of a point
                    if 'sz' in run_props.attrib:
                        font_size = int(run_props.get('sz', 0)) / 100

                    # Text color
                    solid_fill = run_props.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}solidFill')
                    if solid_fill is not None:
                        scheme_color = solid_fill.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}schemeClr')
                        srgb_color = solid_fill.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}srgbClr')

                        if srgb_color is not None:
                            color = '#' + srgb_color.get('val', 'ffffff')
                        elif scheme_color is not None:
                            color_map = {
                                'lt1': '#ffffff',
                                'dk1': '#000000',
                                'accent1': '#0066cc',
                            }
                            color = color_map.get(scheme_color.get('val'), '#000000')

                runs.append({
                    'text': text_content,
                    'bold': bold,
                    'italic': italic,
                    'color': color,
                    'fontSize': font_size
                })

    return runs


def extract_shape_position(shape_elem, slide_width, slide_height):
    """Extract a shape's position and size in slide-relative percentages."""
    position = {
        'x': 0,
        'y': 0,
        'width': 100,
        'height': 20,
    }

    xfrm = shape_elem.find(f'.//{{{DRAW_NS}}}xfrm')
    if xfrm is None:
        return position

    off = xfrm.find(f'.//{{{DRAW_NS}}}off')
    ext = xfrm.find(f'.//{{{DRAW_NS}}}ext')

    if off is not None:
        position['x'] = round(emu_to_percent(off.get('x', 0), slide_width), 2)
        position['y'] = round(emu_to_percent(off.get('y', 0), slide_height), 2)

    if ext is not None:
        position['width'] = round(emu_to_percent(ext.get('cx', 0), slide_width), 2)
        position['height'] = round(emu_to_percent(ext.get('cy', 0), slide_height), 2)

    return position


def extract_text_runs(paragraph_elem):
    runs = []

    for text_run in paragraph_elem.findall(f'.//{{{DRAW_NS}}}r'):
        text_elem = text_run.find(f'.//{{{DRAW_NS}}}t')
        if text_elem is None:
            continue

        text_content = text_elem.text or ''
        if not text_content.strip():
            continue

        run_props = text_run.find(f'.//{{{DRAW_NS}}}rPr')
        bold = False
        italic = False
        color = '#000000'
        font_size = 18

        if run_props is not None:
            bold = run_props.get('b') == '1'
            italic = run_props.get('i') == '1'

            if 'sz' in run_props.attrib:
                try:
                    font_size = int(run_props.get('sz', 0)) / 100
                except Exception:
                    font_size = 18

            solid_fill = run_props.find(f'.//{{{DRAW_NS}}}solidFill')
            if solid_fill is not None:
                srgb_color = solid_fill.find(f'.//{{{DRAW_NS}}}srgbClr')
                if srgb_color is not None:
                    color = '#' + srgb_color.get('val', '000000')

        runs.append({
            'text': text_content,
            'bold': bold,
            'italic': italic,
            'color': color,
            'fontSize': font_size,
        })

    return runs

def parse_pptx(file_path):
    """Parse PPTX file and extract slides with formatted content"""
    slides = []

    try:
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            # Read slide dimensions from presentation.xml
            slide_width = 9144000
            slide_height = 5143500

            try:
                pres_xml = zip_ref.read('ppt/presentation.xml').decode('utf-8')
                pres_root = ET.fromstring(pres_xml)
                sld_sz = pres_root.find(f'.//{{{PML_NS}}}sldSz')
                if sld_sz is not None:
                    cx_val = sld_sz.get('cx')
                    cy_val = sld_sz.get('cy')
                    if cx_val and cy_val:
                        try:
                            slide_width = int(cx_val)
                            slide_height = int(cy_val)
                            print(f"Found slide dimensions: {slide_width} x {slide_height}")
                        except ValueError as e:
                            print(f"Could not parse slide dimensions: {e}")
            except Exception as e:
                print(f"Could not read presentation dimensions: {e}")
                pass

            # Get all slide files, being careful to ignore files like slideLayouts or slideProperties
            slide_files = [f for f in zip_ref.namelist() if f.startswith('ppt/slides/slide') and f.endswith('.xml')]
            # Only keep files that are slideN.xml where N is a number
            def get_slide_num(f):
                name = os.path.basename(f) # slide1.xml
                num_part = name.replace('slide', '').replace('.xml', '')
                try:
                    return int(num_part)
                except ValueError:
                    return None

            slide_files = [f for f in slide_files if get_slide_num(f) is not None]
            slide_files.sort(key=get_slide_num)

            for slide_file in slide_files:
                try:
                    slide_num_value = get_slide_num(slide_file)
                    if slide_num_value is None:
                        continue
                    slide_num = str(slide_num_value)

                    # Read slide XML
                    slide_xml = zip_ref.read(slide_file).decode('utf-8')
                    slide_root = ET.fromstring(slide_xml)

                    # Read slide relationships
                    rel_file = f'ppt/slides/_rels/slide{slide_num}.xml.rels'
                    relationships = {}
                    try:
                        rel_xml = zip_ref.read(rel_file).decode('utf-8')
                        rel_root = ET.fromstring(rel_xml)
                        for rel in rel_root.findall(f'.//{{{PKG_REL_NS}}}Relationship'):
                            rel_id = rel.get('Id')
                            target = rel.get('Target')
                            relationships[rel_id] = target
                    except:
                        pass

                    # Extract text elements and positioned text boxes
                    text_elements = []
                    text_boxes = []
                    full_text = ''

                    for shape in slide_root.findall(f'.//{{{PML_NS}}}sp'):
                        text_body = shape.find(f'.//{{{PML_NS}}}txBody')
                        if text_body is None:
                            continue

                        position = extract_shape_position(shape, slide_width, slide_height)
                        paragraphs = text_body.findall(f'.//{{{DRAW_NS}}}p')

                        box_runs = []
                        alignment = 'l'
                        is_bullet = False
                        level = 0

                        is_title_shape = False
                        nv_props = shape.find(f'.//{{{PML_NS}}}nvSpPr')
                        if nv_props is not None:
                            ph = nv_props.find(f'.//{{{PML_NS}}}ph')
                            if ph is not None:
                                ph_type = ph.get('type', '')
                                is_title_shape = ph_type in ('title', 'ctrTitle', 'subTitle')

                        for para_index, para in enumerate(paragraphs):
                            runs = extract_text_runs(para)
                            if not runs:
                                continue

                            para_text = ''.join([r['text'] for r in runs])
                            if para_text.strip():
                                full_text += para_text + '\n'

                            ppr = para.find(f'.//{{{DRAW_NS}}}pPr')
                            if ppr is not None:
                                align_val = ppr.get('algn')
                                if align_val:
                                    alignment = align_val

                                if ppr.find(f'.//{{{DRAW_NS}}}buChar') is not None or ppr.find(f'.//{{{DRAW_NS}}}buFont') is not None:
                                    is_bullet = True

                                lvl_val = ppr.get('lvl')
                                if lvl_val is not None:
                                    try:
                                        level = int(lvl_val)
                                    except Exception:
                                        level = 0

                            box_runs.extend(runs)

                        if box_runs:
                            text_boxes.append({
                                'runs': box_runs,
                                'type': 'title' if is_title_shape and len(text_boxes) == 0 else 'body',
                                'level': level if is_bullet else None,
                                'isBullet': is_bullet,
                                'alignment': alignment,
                                'x': position['x'],
                                'y': position['y'],
                                'width': position['width'],
                                'height': position['height'],
                            })

                            text_elements.append({
                                'runs': box_runs,
                                'type': 'title' if is_title_shape and len(text_elements) == 0 else 'body',
                                'level': level if is_bullet else None,
                                'isBullet': is_bullet,
                                'alignment': alignment,
                            })

                    # Extract images
                    images = []
                    for pic in slide_root.findall('.//{http://schemas.openxmlformats.org/presentationml/2006/main}pic'):
                        # Get relationship ID
                        blip = pic.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                        if blip is not None:
                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')

                            if embed and embed in relationships:
                                image_path = relationships[embed]
                                if image_path.startswith('../'):
                                    image_path = image_path[3:]
                                image_path = f'ppt/{image_path}'

                                try:
                                    image_data = zip_ref.read(image_path)
                                    image_base64 = base64.b64encode(image_data).decode('utf-8')

                                    # Determine image format
                                    ext = Path(image_path).suffix.lower()
                                    mime_types = {
                                        '.png': 'png',
                                        '.jpg': 'jpeg',
                                        '.jpeg': 'jpeg',
                                        '.gif': 'gif',
                                        '.bmp': 'bmp',
                                    }
                                    mime = mime_types.get(ext, 'png')

                                    # Get position and size from xfrm
                                    xfrm = pic.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm')
                                    x, y, width, height = 0, 0, 15, 15

                                    if xfrm is not None:
                                        off = xfrm.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}off')
                                        ext_elem = xfrm.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ext')

                                        if off is not None:
                                            x = (int(off.get('x', 0)) / slide_width) * 100
                                            y = (int(off.get('y', 0)) / slide_height) * 100

                                        if ext_elem is not None:
                                            width = (int(ext_elem.get('cx', 1000000)) / slide_width) * 100
                                            height = (int(ext_elem.get('cy', 1000000)) / slide_height) * 100

                                    images.append({
                                        'id': embed,
                                        'data': f'data:image/{mime};base64,{image_base64}',
                                        'x': round(x, 2),
                                        'y': round(y, 2),
                                        'width': round(width, 2),
                                        'height': round(height, 2)
                                    })
                                except Exception as e:
                                    print(f"Error loading image: {e}")

                    # Get slide title
                    title = 'Slide ' + slide_num
                    if text_elements and text_elements[0]['runs']:
                        title = ''.join([r['text'] for r in text_elements[0]['runs']])[:50]

                    slides.append({
                        'id': f'slide-{slide_num}',
                        'number': int(slide_num),
                        'title': title,
                        'textElements': text_elements,
                        'textBoxes': text_boxes,
                        'images': images,
                        'fullText': full_text,
                        'backgroundColor': '#ffffff',
                        'width': slide_width,
                        'height': slide_height
                    })

                except Exception as e:
                    print(f"Error parsing slide {slide_file}: {e}")
                    continue

        return slides

    except Exception as e:
        print(f"Error parsing PPTX: {e}")
        return []

@app.route('/api/upload-pptx', methods=['POST'])
def upload_pptx():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Only .pptx files are allowed'}), 400

        # Save temporarily using system temp directory
        filename = secure_filename(file.filename)
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, filename)

        print(f"Saving file to: {temp_path}")
        file.save(temp_path)

        # Verify file was saved
        if not os.path.exists(temp_path):
            return jsonify({'error': 'Failed to save uploaded file'}), 500

        file_size = os.path.getsize(temp_path)
        print(f"File saved, size: {file_size} bytes")

        render_mode = request.form.get('renderMode', 'pixel').lower()
        if render_mode == 'editable':
            slides = parse_pptx(temp_path)
        else:
            # Default: pixel-perfect rendering
            slides = convert_pptx_to_png_slides(temp_path)

        # Clean up
        try:
            os.remove(temp_path)
        except:
            pass

        return jsonify({
            'success': True,
            'slides': slides,
            'total': len(slides)
        }), 200

    except Exception as e:
        print(f"Upload error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/pdf-to-word', methods=['POST'])
def pdf_to_word():
    """
    Convert PDF to DOCX format with high-fidelity preservation of formatting.

    Request:
        - file (multipart/form-data): PDF file to convert

    Response (200):
        {
            "success": true,
            "docxBase64": "base64_encoded_docx_bytes",
            "docxFilename": "document.docx",
            "metadata": {
                "originalSize": 12345,
                "convertedSize": 54321,
                "pages": 5,
                "processTime": 2.34
            }
        }

    Response (400/500):
        {
            "success": false,
            "error": "Error message",
            "errorCode": "ERROR_CODE"
        }
    """
    import time
    start_time = time.time()

    try:
        # Validate request
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'No file provided in request',
                'errorCode': 'NO_FILE'
            }), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'No file selected',
                'errorCode': 'EMPTY_FILENAME'
            }), 400

        # Validate file extension
        ext = Path(file.filename).suffix.lower()
        if ext != '.pdf':
            return jsonify({
                'success': False,
                'error': f'Invalid file type: {ext}. Only .pdf files are supported.',
                'errorCode': 'INVALID_FILE_TYPE'
            }), 400

        # Validate file size (50MB max)
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        if file_size > 50 * 1024 * 1024:
            return jsonify({
                'success': False,
                'error': 'File size exceeds 50MB limit',
                'errorCode': 'FILE_TOO_LARGE'
            }), 400

        if file_size == 0:
            return jsonify({
                'success': False,
                'error': 'File is empty',
                'errorCode': 'EMPTY_FILE'
            }), 400

        filename = secure_filename(file.filename)
        temp_dir = tempfile.gettempdir()
        temp_pdf_path = os.path.join(temp_dir, filename)

        # Save uploaded file
        try:
            file.save(temp_pdf_path)
        except Exception as save_error:
            return jsonify({
                'success': False,
                'error': f'Failed to save uploaded file: {str(save_error)}',
                'errorCode': 'SAVE_ERROR'
            }), 500

        if not os.path.exists(temp_pdf_path):
            return jsonify({
                'success': False,
                'error': 'Failed to save uploaded PDF',
                'errorCode': 'SAVE_VERIFY_ERROR'
            }), 500

        try:
            # Perform conversion
            docx_base64 = convert_pdf_to_docx(temp_pdf_path)

            # Calculate conversion metrics
            process_time = time.time() - start_time

            # Decode to get actual size
            docx_bytes = base64.b64decode(docx_base64)

            # Extract page count from PDF
            try:
                pdf_doc = fitz.open(temp_pdf_path)
                page_count = len(pdf_doc)
                pdf_doc.close()
            except Exception:
                page_count = 0

            response_data = {
                'success': True,
                'docxBase64': docx_base64,
                'docxFilename': Path(file.filename).stem + '.docx',
                'metadata': {
                    'originalSize': file_size,
                    'convertedSize': len(docx_bytes),
                    'pages': page_count,
                    'processTime': round(process_time, 2)
                }
            }

            return jsonify(response_data), 200

        except RuntimeError as runtime_error:
            print(f"PDF->Word conversion runtime error: {runtime_error}")
            return jsonify({
                'success': False,
                'error': f'Conversion failed: {str(runtime_error)}',
                'errorCode': 'CONVERSION_ERROR'
            }), 500
        except Exception as conversion_error:
            print(f"PDF->Word conversion error: {conversion_error}")
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False,
                'error': 'An unexpected error occurred during conversion. Check server logs for details.',
                'errorCode': 'UNEXPECTED_ERROR'
            }), 500
        finally:
            # Clean up temporary PDF
            try:
                if os.path.exists(temp_pdf_path):
                    os.remove(temp_pdf_path)
            except Exception as cleanup_error:
                print(f"Warning: Could not clean up temp file {temp_pdf_path}: {cleanup_error}")
    except Exception as e:
        print(f"PDF->Word request handler error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': 'Server error processing request',
            'errorCode': 'REQUEST_ERROR'
        }), 500


@app.route('/api/word-to-pdf', methods=['POST'])
def word_to_pdf():
    """Convert Word (DOC/DOCX) to PDF using Cloudmersive."""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        ext = Path(file.filename).suffix.lower()
        if ext not in ['.doc', '.docx']:
            return jsonify({'success': False, 'error': f'Unsupported extension: {ext}'}), 400

        filename = secure_filename(file.filename)
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, filename)
        file.save(temp_path)

        try:
            print(f"Converting {filename} to PDF via Cloudmersive")
            if ext == '.doc':
                api_response = convert_api.convert_document_doc_to_pdf(temp_path)
            else:
                api_response = convert_api.convert_document_docx_to_pdf(temp_path)

            if isinstance(api_response, bytes):
                pdf_base64 = base64.b64encode(api_response).decode('utf-8')
            elif isinstance(api_response, str) and os.path.exists(api_response):
                with open(api_response, 'rb') as f:
                    pdf_base64 = base64.b64encode(f.read()).decode('utf-8')
            else:
                raise RuntimeError("Cloudmersive returned unexpected response type")

            return jsonify({
                'success': True,
                'pdfBase64': pdf_base64,
                'pdfFilename': Path(file.filename).stem + '.pdf'
            }), 200

        except ApiException as e:
            print(f"Cloudmersive API Exception: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except Exception as e:
        print(f"Word-to-PDF error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/pdf-to-pptx', methods=['POST'])
def pdf_to_pptx():
    """Convert PDF to PPTX using Cloudmersive with high-fidelity design preservation."""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        ext = Path(file.filename).suffix.lower()
        if ext != '.pdf':
            return jsonify({'success': False, 'error': f'Unsupported extension: {ext}. Only .pdf is supported.'}), 400

        filename = secure_filename(file.filename)
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, filename)
        file.save(temp_path)

        try:
            print(f"Converting {filename} to PPTX via Cloudmersive")
            api_response = convert_api.convert_document_pdf_to_pptx(temp_path)

            if isinstance(api_response, bytes):
                pptx_base64 = base64.b64encode(api_response).decode('utf-8')
            elif isinstance(api_response, str) and os.path.exists(api_response):
                with open(api_response, 'rb') as f:
                    pptx_base64 = base64.b64encode(f.read()).decode('utf-8')
            else:
                raise RuntimeError("Cloudmersive returned unexpected response type")

            return jsonify({
                'success': True,
                'pptxBase64': pptx_base64,
                'pptxFilename': Path(file.filename).stem + '.pptx'
            }), 200

        except ApiException as e:
            print(f"Cloudmersive API Exception: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except Exception as e:
        print(f"PDF-to-PPTX error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/pptx-to-word', methods=['POST'])
def pptx_to_word():
    """Convert PPTX to DOCX using Cloudmersive with high-fidelity design preservation."""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        ext = Path(file.filename).suffix.lower()
        if ext != '.pptx':
            return jsonify({'success': False, 'error': f'Unsupported extension: {ext}. Only .pptx is supported.'}), 400

        filename = secure_filename(file.filename)
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, filename)
        file.save(temp_path)

        try:
            print(f"Converting {filename} to Word via Cloudmersive (2-step: PPTX->PDF->Word)")

            # Step 1: PPTX to PDF
            pdf_bytes = convert_api.convert_document_pptx_to_pdf(temp_path)

            # Save intermediate PDF to temp file
            temp_pdf_path = os.path.join(temp_dir, f"temp_{filename}.pdf")
            with open(temp_pdf_path, 'wb') as f:
                f.write(pdf_bytes)

            try:
                # Step 2: PDF to DOCX
                api_response = convert_api.convert_document_pdf_to_docx(temp_pdf_path)

                if isinstance(api_response, bytes):
                    docx_base64 = base64.b64encode(api_response).decode('utf-8')
                elif isinstance(api_response, str) and os.path.exists(api_response):
                    with open(api_response, 'rb') as f:
                        docx_base64 = base64.b64encode(f.read()).decode('utf-8')
                else:
                    raise RuntimeError("Cloudmersive returned unexpected response type")

                return jsonify({
                    'success': True,
                    'docxBase64': docx_base64,
                    'docxFilename': Path(file.filename).stem + '.docx'
                }), 200
            finally:
                if os.path.exists(temp_pdf_path):
                    os.remove(temp_pdf_path)

        except ApiException as e:
            print(f"Cloudmersive API Exception: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except Exception as e:
        print(f"PPTX-to-Word error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'}), 200


def sanitize_user_for_response(user):
    serialized = serialize_mongo_document(user)
    if isinstance(serialized, dict):
        serialized.pop('passwordHash', None)
    return serialized


def create_auth_session(user):
    return {
        'token': secrets.token_urlsafe(32),
        'user': sanitize_user_for_response(user),
    }


@app.route('/api/auth/register', methods=['POST'])
def register_account():
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    data = request.get_json(silent=True) or {}
    email = normalize_email(data.get('email') or data.get('editorEmail'))
    display_name = str(data.get('displayName') or data.get('editorName') or '').strip()[:160]
    password = str(data.get('password') or '')

    if not display_name:
        return jsonify({'success': False, 'error': 'Name is required'}), 400
    if not email:
        return jsonify({'success': False, 'error': 'A valid email is required'}), 400
    if len(password) < 6:
        return jsonify({'success': False, 'error': 'Password must be at least 6 characters'}), 400

    existing = db.users.find_one({'email': email})
    if existing and existing.get('passwordHash'):
        return jsonify({'success': False, 'error': 'An account with this email already exists'}), 409

    now = datetime.now(timezone.utc)
    user_doc = {
        'userId': email[:120],
        'displayName': display_name,
        'email': email,
        'passwordHash': generate_password_hash(password),
        'updatedAt': now,
        'lastSeenAt': now,
    }
    db.users.update_one(
        {'email': email},
        {
            '$set': user_doc,
            '$setOnInsert': {'createdAt': now},
        },
        upsert=True,
    )
    user = db.users.find_one({'email': email})
    return jsonify({'success': True, **create_auth_session(user)}), 201


@app.route('/api/auth/login', methods=['POST'])
def login_account():
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    data = request.get_json(silent=True) or {}
    email = normalize_email(data.get('email') or data.get('editorEmail'))
    password = str(data.get('password') or '')
    if not email or not password:
        return jsonify({'success': False, 'error': 'Email and password are required'}), 400

    user = db.users.find_one({'email': email})
    if not user or not user.get('passwordHash') or not check_password_hash(user['passwordHash'], password):
        return jsonify({'success': False, 'error': 'Invalid email or password'}), 401

    now = datetime.now(timezone.utc)
    db.users.update_one({'email': email}, {'$set': {'lastSeenAt': now, 'updatedAt': now}})
    user = db.users.find_one({'email': email})
    return jsonify({'success': True, **create_auth_session(user)}), 200


@app.route('/api/users', methods=['POST'])
def create_user():
    """Create or update an editor user in MongoDB."""
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    data = request.get_json(silent=True) or {}
    actor = get_actor_from_payload(data)
    raw_email = data.get('email') or data.get('userEmail') or data.get('editorEmail')
    if raw_email and not actor.get('email'):
        return jsonify({'success': False, 'error': 'A valid email is required'}), 400
    if not actor['displayName']:
        return jsonify({'success': False, 'error': 'displayName is required'}), 400

    ensure_user(db, actor)
    user = db.users.find_one({'userId': actor['userId']})
    return jsonify({
        'success': True,
        'user': sanitize_user_for_response(user),
    }), 200


@app.route('/api/users', methods=['GET'])
def list_users():
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    try:
        limit = min(max(int(request.args.get('limit', 100)), 1), 500)
    except Exception:
        limit = 100

    cursor = db.users.find().sort('updatedAt', -1).limit(limit)
    return jsonify({
        'success': True,
        'users': [sanitize_user_for_response(user) for user in cursor],
    }), 200


@app.route('/api/files', methods=['POST'])
def create_file_record():
    """Create or update the MongoDB record for an uploaded/opened file."""
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    data = request.get_json(silent=True) or {}
    file_record, validation_error = build_file_record(data)
    if validation_error:
        return jsonify({'success': False, 'error': validation_error}), 400

    actor = file_record['uploadedBy']
    ensure_user(db, actor)

    file_id = file_record['fileId']
    existing = db.files.find_one({'fileId': file_id})
    now = datetime.now(timezone.utc)
    try:
        content_update = save_file_content_to_gridfs(db, file_id, data, existing)
    except (ValueError, RuntimeError) as error:
        return jsonify({'success': False, 'error': str(error)}), 400

    if existing:
        update = {
            'fileName': file_record['fileName'],
            'fileType': file_record['fileType'],
            'originalType': file_record['originalType'],
            'workflow': file_record['workflow'],
            'size': file_record['size'],
            'updatedAt': now,
            **content_update,
        }
        db.files.update_one({'fileId': file_id}, {'$set': update})
    else:
        file_record.update(content_update)
        db.files.insert_one(file_record)

    stored_file = db.files.find_one({'fileId': file_id})
    return jsonify({
        'success': True,
        'file': serialize_mongo_document(stored_file),
    }), 201 if not existing else 200


@app.route('/api/files/<file_id>/content', methods=['PUT'])
def update_file_content(file_id):
    """Save updated file content to GridFS so all users can view the latest version."""
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    file_record = db.files.find_one({'fileId': file_id})
    if file_record is None:
        return jsonify({'success': False, 'error': 'File not found'}), 404

    data = request.get_json(silent=True) or {}
    content_base64 = data.get('contentBase64')
    if not content_base64:
        return jsonify({'success': False, 'error': 'contentBase64 is required'}), 400

    try:
        content_update = save_file_content_to_gridfs(db, file_id, data, file_record)
    except (ValueError, RuntimeError) as error:
        return jsonify({'success': False, 'error': str(error)}), 400

    actor = get_actor_from_payload(data)
    ensure_user(db, actor)
    now = datetime.now(timezone.utc)

    db.files.update_one(
        {'fileId': file_id},
        {
            '$set': {
                **content_update,
                'updatedAt': now,
                'edited': True,
                'lastEditedAt': now,
                'lastEditedBy': actor,
            },
            '$inc': {'editCount': 1},
        },
    )

    stored_file = db.files.find_one({'fileId': file_id})
    return jsonify({
        'success': True,
        'file': serialize_mongo_document(stored_file),
    }), 200


@app.route('/api/files/<file_id>/sender-share', methods=['GET'])
def get_sender_access(file_id):
    """Return the accessToken for the file's original uploader so they can open it via shared URL."""
    db = get_mongo_db()
    if db is None:
        return jsonify({'success': False, 'error': 'MongoDB is not available'}), 503

    file_record = db.files.find_one({'fileId': file_id})
    if file_record is None:
        return jsonify({'success': False, 'error': 'File not found'}), 404

    # Find any share record for this file (sender's own token)
    share = db.file_shares.find_one({'fileId': file_id}, sort=[('createdAt', 1)])
    if share:
        return jsonify({
            'success': True,
            'accessToken': share.get('accessToken'),
            'accessUrl': share.get('accessUrl'),
        }), 200

    # No share exists yet — create a self-access token for the uploader
    uploaded_by = file_record.get('uploadedBy', {})
    if not uploaded_by.get('userId'):
        return jsonify({'success': False, 'error': 'No uploader info found'}), 404

    now = datetime.now(timezone.utc)
    access_token = secrets.token_urlsafe(32)
    access_url = f'{get_frontend_public_url()}/shared/{access_token}'
    self_share = {
        'fileId': file_id,
        'fileName': file_record.get('fileName', ''),
        'fileType': file_record.get('fileType', ''),
        'permission': 'edit',
        'accessToken': access_token,
        'accessUrl': access_url,
        'emailStatus': {'sent': False, 'configured': False},
        'sharedBy': uploaded_by,
        'sharedWith': uploaded_by,
        'createdAt': now,
        'updatedAt': now,
    }
    db.file_shares.update_one(
        {'fileId': file_id, 'sharedWith.userId': uploaded_by['userId']},
        {'$set': self_share, '$setOnInsert': {'firstSharedAt': now}},
        upsert=True,
    )
    return jsonify({
        'success': True,
        'accessToken': access_token,
        'accessUrl': access_url,
    }), 201


@app.route('/api/files', methods=['GET'])
def list_file_records():
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    try:
        limit = min(max(int(request.args.get('limit', 50)), 1), 200)
    except Exception:
        limit = 50

    cursor = db.files.find().sort('uploadedAt', -1).limit(limit)
    return jsonify({
        'success': True,
        'files': [serialize_mongo_document(file_record) for file_record in cursor],
    }), 200


@app.route('/api/files/<file_id>', methods=['GET'])
def get_file_record(file_id):
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    file_record = db.files.find_one({'fileId': file_id})
    if file_record is None:
        return jsonify({'success': False, 'error': 'File not found'}), 404

    return jsonify({
        'success': True,
        'file': serialize_mongo_document(file_record),
    }), 200


@app.route('/api/files/<file_id>/shares', methods=['GET'])
def get_file_shares(file_id):
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    file_record = db.files.find_one({'fileId': file_id})
    if file_record is None:
        return jsonify({'success': False, 'error': 'File not found'}), 404

    cursor = db.file_shares.find({'fileId': file_id}).sort('createdAt', -1)
    return jsonify({
        'success': True,
        'fileId': file_id,
        'shares': [serialize_mongo_document(share) for share in cursor],
    }), 200


@app.route('/api/users/<path:user_email>/shares', methods=['GET'])
def get_user_share_history(user_email):
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    email = normalize_email(user_email)
    if not email:
        return jsonify({'success': False, 'error': 'A valid user email is required'}), 400

    received_cursor = db.file_shares.find({'sharedWith.email': email}).sort('updatedAt', -1).limit(100)
    sent_cursor = db.file_shares.find({'sharedBy.email': email}).sort('updatedAt', -1).limit(100)
    received = [serialize_mongo_document(share) for share in received_cursor]
    sent = [serialize_mongo_document(share) for share in sent_cursor]

    return jsonify({
        'success': True,
        'email': email,
        'received': received,
        'sent': sent,
        'notificationCount': len(received),
    }), 200


@app.route('/api/files/<file_id>/shares', methods=['POST'])
def share_file(file_id):
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    file_record = db.files.find_one({'fileId': file_id})
    if file_record is None:
        return jsonify({'success': False, 'error': 'File not found'}), 404

    data = request.get_json(silent=True) or {}
    shared_with_email = normalize_email(data.get('sharedWithEmail') or data.get('email') or data.get('sharedWith'))
    shared_with_name = str(data.get('sharedWithName') or data.get('sharedWith') or shared_with_email or '').strip()
    permission = str(data.get('permission') or 'view').strip().lower()

    if not shared_with_email:
        return jsonify({'success': False, 'error': 'A valid sharedWithEmail is required'}), 400
    if permission not in ['view', 'edit']:
        return jsonify({'success': False, 'error': 'permission must be view or edit'}), 400

    now = datetime.now(timezone.utc)
    actor = get_actor_from_payload(data)
    ensure_user(db, actor)
    shared_with = {
        'userId': shared_with_email[:120],
        'displayName': (shared_with_name or shared_with_email)[:160],
        'email': shared_with_email,
    }
    ensure_user(db, shared_with)

    existing_share = db.file_shares.find_one({'fileId': file_id, 'sharedWith.userId': shared_with['userId']})
    access_token = existing_share.get('accessToken') if existing_share else secrets.token_urlsafe(32)
    access_url = f'{get_frontend_public_url()}/shared/{access_token}'
    email_status = send_share_email(
        shared_with_email,
        shared_with['displayName'],
        actor['displayName'],
        file_record.get('fileName', 'file'),
        access_url,
        permission,
    )

    share = {
        'fileId': file_id[:160],
        'fileName': file_record.get('fileName', ''),
        'fileType': file_record.get('fileType', ''),
        'permission': permission,
        'accessToken': access_token,
        'accessUrl': access_url,
        'emailStatus': email_status,
        'sharedBy': actor,
        'sharedWith': shared_with,
        'createdAt': existing_share.get('createdAt') if existing_share else now,
        'updatedAt': now,
    }

    db.file_shares.update_one(
        {'fileId': file_id, 'sharedWith.userId': shared_with['userId']},
        {
            '$set': share,
            '$setOnInsert': {'firstSharedAt': now},
        },
        upsert=True,
    )
    stored_share = db.file_shares.find_one({'fileId': file_id, 'sharedWith.userId': shared_with['userId']})

    event = {
        'fileId': file_id[:160],
        'fileName': file_record.get('fileName', ''),
        'fileType': file_record.get('fileType', ''),
        'action': f'share {permission}',
        'editor': file_record.get('fileType', 'unknown'),
        'userId': actor['userId'],
        'editorName': actor['displayName'],
        'metadata': {
            'sharedWith': shared_with,
            'permission': permission,
            'accessUrl': access_url,
            'emailStatus': email_status,
        },
        'createdAt': now,
        'updatedAt': now,
        'ipAddress': request.remote_addr,
        'userAgent': get_request_user_agent(),
    }
    db.edit_events.insert_one(event)
    db.files.update_one(
        {'fileId': file_id},
        {
            '$set': {
                'shared': True,
                'lastSharedAt': now,
                'lastSharedBy': actor,
                'updatedAt': now,
            },
            '$inc': {'shareCount': 1},
        },
    )

    return jsonify({
        'success': True,
        'share': serialize_mongo_document(stored_share),
        'event': serialize_mongo_document(event),
        'emailStatus': email_status,
        'accessUrl': access_url,
    }), 201


@app.route('/shared/<access_token>', methods=['GET'])
def shared_file_access_page(access_token):
    return redirect(f'{get_frontend_public_url()}/shared/{access_token}', code=302)


@app.route('/api/shared/<access_token>', methods=['GET'])
def get_shared_file(access_token):
    db = get_mongo_db()
    if db is None:
        return jsonify({'success': False, 'error': 'MongoDB is not available'}), 503

    share = db.file_shares.find_one({'accessToken': access_token})
    if share is None:
        return jsonify({'success': False, 'error': 'Shared file link not found'}), 404

    file_record = db.files.find_one({'fileId': share.get('fileId')})
    if file_record is None:
        return jsonify({'success': False, 'error': 'File not found'}), 404

    content_base64 = ''
    if file_record.get('contentGridFsId'):
        try:
            import gridfs
            from bson import ObjectId
            fs = gridfs.GridFS(db)
            grid_file = fs.get(ObjectId(file_record['contentGridFsId']))
            content_base64 = base64.b64encode(grid_file.read()).decode('utf-8')
        except Exception as error:
            return jsonify({'success': False, 'error': f'Could not load file content: {error}'}), 500

    return jsonify({
        'success': True,
        'share': serialize_mongo_document(share),
        'file': serialize_mongo_document(file_record),
        'contentBase64': content_base64,
        'downloadUrl': f'{get_public_base_url()}/api/shared/{access_token}/download',
    }), 200


@app.route('/api/shared/<access_token>/download', methods=['GET'])
def download_shared_file(access_token):
    db = get_mongo_db()
    if db is None:
        return jsonify({'success': False, 'error': 'MongoDB is not available'}), 503

    share = db.file_shares.find_one({'accessToken': access_token})
    if share is None:
        return jsonify({'success': False, 'error': 'Shared file link not found'}), 404

    file_record = db.files.find_one({'fileId': share.get('fileId')})
    if file_record is None:
        return jsonify({'success': False, 'error': 'File not found'}), 404
    if not file_record.get('contentGridFsId'):
        return jsonify({'success': False, 'error': 'File content is not stored for this file'}), 404

    try:
        import gridfs
        from bson import ObjectId
        fs = gridfs.GridFS(db)
        grid_file = fs.get(ObjectId(file_record['contentGridFsId']))
        data = grid_file.read()
    except Exception as error:
        return jsonify({'success': False, 'error': f'Could not load file content: {error}'}), 500

    return send_file(
        BytesIO(data),
        mimetype=file_record.get('contentType') or 'application/octet-stream',
        as_attachment=True,
        download_name=file_record.get('fileName') or 'shared-file',
    )


@app.route('/api/edit-events', methods=['POST'])
def create_edit_event():
    """Record that a user changed or saved a file."""
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    data = request.get_json(silent=True) or {}
    file_id = str(data.get('fileId') or '').strip()
    file_name = str(data.get('fileName') or '').strip()
    action = str(data.get('action') or 'edit').strip()[:80]

    if not file_id or not file_name:
        return jsonify({
            'success': False,
            'error': 'fileId and fileName are required',
        }), 400

    now = datetime.now(timezone.utc)
    actor = get_actor_from_payload(data)
    ensure_user(db, actor)
    metadata = data.get('metadata') if isinstance(data.get('metadata'), dict) else {}
    event = {
        'fileId': file_id[:160],
        'fileName': file_name[:260],
        'fileType': str(data.get('fileType') or '')[:40],
        'action': action,
        'editor': str(data.get('editor') or 'unknown')[:80],
        'userId': actor['userId'],
        'editorName': actor['displayName'],
        'metadata': metadata,
        'createdAt': now,
        'updatedAt': now,
        'ipAddress': request.remote_addr,
        'userAgent': get_request_user_agent(),
    }

    result = db.edit_events.insert_one(event)
    event['_id'] = result.inserted_id

    file_update = {
        'fileId': event['fileId'],
        'fileName': event['fileName'],
        'fileType': event['fileType'],
        'updatedAt': now,
        'edited': True,
        'lastEdit': event,
        'lastEditedAt': now,
        'lastEditedBy': actor,
    }
    db.files.update_one(
        {'fileId': event['fileId']},
        {
            '$set': file_update,
            '$inc': {'editCount': 1},
            '$setOnInsert': {
                'uploadedAt': now,
                'uploadedBy': actor,
                'originalType': '',
                'workflow': '',
                'size': 0,
                'ipAddress': request.remote_addr,
                'userAgent': get_request_user_agent(),
            },
        },
        upsert=True,
    )

    return jsonify({
        'success': True,
        'event': serialize_mongo_document(event),
    }), 201


@app.route('/api/files/<file_id>/edits', methods=['GET'])
def get_file_edit_events(file_id):
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    try:
        limit = min(max(int(request.args.get('limit', 50)), 1), 200)
    except Exception:
        limit = 50

    cursor = db.edit_events.find({'fileId': file_id}).sort('createdAt', -1).limit(limit)
    events = [serialize_mongo_document(event) for event in cursor]
    file_record = db.files.find_one({'fileId': file_id})

    return jsonify({
        'success': True,
        'fileId': file_id,
        'file': serialize_mongo_document(file_record),
        'edited': len(events) > 0,
        'events': events,
        'lastEdit': events[0] if events else None,
    }), 200


@app.route('/api/files/<file_id>/edit-status', methods=['GET'])
def get_file_edit_status(file_id):
    db = get_mongo_db()
    if db is None:
        return jsonify({
            'success': False,
            'error': 'MongoDB is not available',
            'details': _mongo_error,
        }), 503

    file_record = db.files.find_one({'fileId': file_id})
    last_event = db.edit_events.find_one({'fileId': file_id}, sort=[('createdAt', -1)])

    return jsonify({
        'success': True,
        'fileId': file_id,
        'file': serialize_mongo_document(file_record),
        'edited': last_event is not None,
        'lastEdit': serialize_mongo_document(last_event),
    }), 200

if __name__ == '__main__':
    app.run(debug=True, port=5000)
